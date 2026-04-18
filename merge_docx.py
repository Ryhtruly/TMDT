import os
from docx import Document
from docxcompose.composer import Composer

def main():
    report1_path = "Website quản lý kho vận - Report 1 - Research thực tế.docx"
    database_path = "Database (1) - DA SUA.docx"
    output_path = "Bao_Cao_Tong_Hop_Hoan_Chinh.docx"

    print("Opening master document...")
    master = Document(report1_path)
    composer = Composer(master)

    print("Appending Database document...")
    try:
        doc2 = Document(database_path)
        # Add a page break before appending
        master.add_page_break()
        composer.append(doc2)
    except Exception as e:
        print(f"Error appending database doc: {e}")

    # Now add the tech spec
    doc = composer.doc
    doc.add_page_break()
    
    doc.add_heading("PHẦN TIẾP THEO: REPORT 2 & CẤU TRÚC CÔNG NGHỆ (TECH SPEC)", level=1)
    
    doc.add_heading("1. Kiến Trúc Tổng Thể Hệ Thống (Monorepo Architecture)", level=2)
    doc.add_paragraph("Hệ thống được thiết kế theo kiến trúc Monorepo, chia làm 4 module chính chạy độc lập nhưng có sự kế thừa mã nguồn và dữ liệu chặt chẽ:")
    
    p1 = doc.add_paragraph()
    p1.add_run("- Webservice (Backend API): ").bold = True
    p1.add_run("Xây dựng trên nền tảng NodeJS, ExpressJS 5, và TypeScript. Chịu trách nhiệm cung cấp RESTful APIs cho toàn bộ các Client. Sử dụng PostgreSQL với connection pooling.")
    
    p2 = doc.add_paragraph()
    p2.add_run("- Client Web (Phân hệ Shop): ").bold = True
    p2.add_run("Ứng dụng dành cho nhà bán hàng (Người gửi). Xây dựng bằng React 19, Vite. Cho phép tạo đơn hàng, quản lý dòng tiền thu hộ (COD) và tra cứu lộ trình chuyển phát (Tracking) theo thời gian thực.")
    
    p3 = doc.add_paragraph()
    p3.add_run("- Admin Web (Phân hệ Quản trị): ").bold = True
    p3.add_run("Bảng điều khiển (Dashboard) dành cho người quản lý cấp cao. Tích hợp thư viện Recharts để thống kê doanh thu và react-force-graph-3d để vẽ bản đồ mạng lưới Hub-Spoke (Kho trung chuyển và Cửa hàng) không gian 3D tương tác.")
    
    p4 = doc.add_paragraph()
    p4.add_run("- Mobile Web (Phân hệ Shipper & Stockkeeper): ").bold = True
    p4.add_run("Ứng dụng PWA dành cho nhân viên lấy/giao hàng và thủ kho hiện trường. Tối ưu hóa UI/UX trên mobile. Tích hợp html5-qrcode giúp quét mã vận đơn (Barcode) trực tiếp qua Camera điện thoại hoặc URL Webcam cắm ngoài.")

    doc.add_heading("2. Đặc Tả Cơ Sở Dữ Liệu PostgreSQL Mở Rộng", level=2)
    doc.add_paragraph("Như đã phân tích ở phần Database, hệ thống xử lý tính vẹn toàn dữ liệu thông qua cơ chế Transaction của PostgreSQL (BEGIN ... COMMIT ... ROLLBACK). Các kỹ thuật khóa bi quan (Pessimistic Locking - FOR UPDATE) được sử dụng khi hệ thống thực hiện thao tác rút tiền hoặc nạp tiền ví, nhằm chống lại lỗi bất đồng bộ (Race Condition).")
    
    doc.add_heading("3. Các Luồng Nghiệp Vụ Cốt Lõi (Core Workflows)", level=2)
    
    doc.add_heading("3.1 Luồng Đối Soát Dòng Tiền & COD", level=3)
    doc.add_paragraph("Khi một kiện hàng được Giao Thành Công, Shipper sẽ cập nhật trạng thái trên Mobile App. Hệ thống lập tức ghi nhận tiền thu hộ (COD) và đồng thời thực hiện cập nhật đối soát cho Wallet (Ví ảo) của Shop. Nhờ vậy Shop có thể thấy số tiền về ví theo thời gian thực và tạo yêu cầu Rút Số Dư về tài khoản Ngân Hàng bất kỳ lúc nào. Luồng nạp tiền tự động được tự động hóa qua quét mã QR với Webhook liên kết từ PayOS.")

    doc.add_heading("3.2 Luồng Định Tuyến Giao Hàng (Routing Mechanism)", level=3)
    doc.add_paragraph("Logistics nội bộ sử dụng mạng lưới Hub/Spoke. Khi Shop lên đơn chốt điểm Giao và điểm Nhận, Backend sẽ mapping địa chỉ của 2 điểm này với các Spoke (Bưu cục phát) và Hub (Kho cấp 1) tương ứng. Nếu cùng nội thành, hàng hoá chỉ qua 1 Hub. Nếu khác vùng, hàng hoá sẽ được điều vận qua Xe tải liên tỉnh giữa 2 Hub lớn trước khi đổ bộ về Spoke phát. Cước phí (Pricing Rules) sẽ được tự động nội suy qua nấc trọng lượng và Phân loại vùng (Province tới District).")

    doc.add_heading("4. Tính Năng Bảo Mật Hệ Thống", level=2)
    doc.add_paragraph("Toàn bộ các end-points quan trọng đều được ủy quyền bảo mật thông qua chuẩn Bearer JWT (JSON Web Tokens). Phân chia quyền hạn sâu đến 4 phân lớp: ADMIN, SHOP, SHIPPER, STOCKKEEPER. Ví dụ một Shipper chỉ có thể quét nhận vận đơn khi Shipper đó đang phụ thuộc vào Spoke đó và thuộc đúng ca làm việc đó. Mật khẩu trên database được encrypt 1 chiều chống dịch ngược bằng các vòng lặp Hash của thuật toán Bcrypt.")

    print(f"Saving merged document to {output_path}...")
    doc.save(output_path)
    print("Done!")

if __name__ == "__main__":
    main()
